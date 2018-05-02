#!/usr/bin/python
# -*- coding: utf-8 -*-
"""
Usage: python create_docs.py

    This script have not parameters, it reads data.pickle and creates 3 files:
    reduced_names.docx with reduced names;
    full_names.docx with full names;
    reduced_names_and_ids.docx with reduced names and unic ids;
    full_names_and_ids.docx with full names and unic ids.
    ids in the two last files matches for equal comments.

"""
import pickle
import os
from docopt import docopt
from docx import Document
from docx.shared import Inches

arguments = docopt(__doc__)


document_with_reduced_names = Document()
document_without_reduced_names = Document()
document_with_reduced_names_and_ids = Document()
document_without_reduced_names_and_ids = Document()

with open('data.pickle', 'rb') as f:
    comments = pickle.load(f)

def reduce_name(name):
    new_name =  u''.join( [ s[0] for s in name.split(',')[0].split()[:3] ] ).upper()
    new_name = u''.join([ n for n in new_name if n.isupper() ])
    return new_name

for inx in xrange(len(comments)):
    name, text, time_stamp = comments[inx]

    reduced_name = reduce_name(name)
    document_with_reduced_names.add_paragraph(u'{}: {}'.format(reduced_name, text))
    document_without_reduced_names.add_paragraph(u'{}: {}'.format(name, text))
    document_with_reduced_names_and_ids.add_paragraph(u'{}, #{}: {}'.format(reduced_name, inx, text))
    document_without_reduced_names_and_ids.add_paragraph(u'{}, #{}: {}'.format(name, inx, text))


basepath = os.path.dirname(__file__)

document_with_reduced_names.save(os.path.abspath(os.path.join(basepath, 'reduced_names.docx')))
document_without_reduced_names.save(os.path.abspath(os.path.join(basepath, 'full_names.docx')))
document_with_reduced_names_and_ids.save(os.path.abspath(os.path.join(basepath, 'reduced_names_and_ids.docx')))
document_without_reduced_names_and_ids.save(os.path.abspath(os.path.join(basepath, 'full_names_and_ids.docx')))
